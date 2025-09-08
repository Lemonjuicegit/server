from docxtpl import DocxTemplate
import docxtpl
import os
import re
import geopandas as gpd
import pandas as pd
from docx import Document
from docxcompose.composer import Composer
from docx.shared import Mm
template_path = r'E:\exploitation\webpython\routers\大足占补平衡申报\template\低效林园地开发项目竣工前后影像对比.docx'

def generate_opinion(dkxh,jgq_img_path='',jgh_img_path=''):
    doc = DocxTemplate(template_path)
    jgq_image = docxtpl.InlineImage(doc, jgq_img_path, height=Mm(100)) if jgq_img_path else jgq_img_path
    jgh_image = docxtpl.InlineImage(doc, jgh_img_path, height=Mm(100)) if jgh_img_path else jgh_img_path
    context = {
        'dkxh':dkxh,
        'jgq_img':jgq_image,
        'jgh_img':jgh_image,
    }
    doc.render(context)
    return doc

def generate_opinion_all(jgq_img_dir,jgh_img_dir,sava):
    # gdf = gpd.read_file(gdb,layer="南大街")
    composer = Composer(Document())
    docxlist = []
    jgq_path = os.listdir(jgq_img_dir)
    jgh_path = os.listdir(jgh_img_dir)
    jgq_dkxh_dict = {re.search(r'地块-[0-9]*',v).group():v for v in jgq_path}
    jgh_dkxh_dict = {re.search(r'地块-[0-9]*',v).group():v for v in jgh_path}
    jgq_dkxh_set = {re.search(r'地块-[0-9]*',v).group() for v in jgq_path}
    jgh_dkxh_set = {re.search(r'地块-[0-9]*',v).group() for v in jgh_path}
    jgq_ = jgq_dkxh_set - jgh_dkxh_set
    jgh_ = jgh_dkxh_set - jgq_dkxh_set
    jgq_jgh = jgq_dkxh_set & jgh_dkxh_set

    for jgq in jgq_:
        doc = generate_opinion(jgq,os.path.join(jgq_img_dir,jgq_dkxh_dict[jgq]))
        docxlist.append(doc)
    for jgq in jgh_:
        doc = generate_opinion(jgq,jgh_img_path=os.path.join(jgh_img_dir,jgh_dkxh_dict[jgq]))
        docxlist.append(doc)
    for jgq in jgq_jgh:
        doc = generate_opinion(jgq,os.path.join(jgq_img_dir,jgq_dkxh_dict[jgq]),os.path.join(jgh_img_dir,jgh_dkxh_dict[jgq]))
        docxlist.append(doc)

    for docx in docxlist:
        composer.append(docx)
    composer.save(sava)
 
def xg(doc_path,dlpath):
    dl_df = pd.read_excel(dlpath)
    doc = Document(doc_path)
    for table in doc.tables:
        text = table.cell(0, 0).paragraphs[0].text
        dl = dl_df[dl_df["编号"] == text]
        if len(dl) == 0:
            with open(r"C:\Users\Administrator\Desktop\新建文件夹\SQ\err.txt", "a",encoding="utf-8") as f:
                f.write(text)
                f.write("\n")
            continue
        else:
            dl = dl.values[0]
        table.cell(0, 0).paragraphs[0].runs[0].text = f"{text}（实施前地类：{dl[0]}）"
    doc.save(r"C:\Users\Administrator\Desktop\新建文件夹\GL\大足区国梁镇等（27）个镇低效林园地开发项目竣工前后影像对比核实后-地类.docx")
if __name__ == "__main__":
    datapath = r"C:\Users\Administrator\Desktop\新建文件夹\GL\大足区国梁镇等（27）个镇低效林园地开发项目竣工前后影像对比核实后.docx"
    dlpath = r"C:\Users\Administrator\Desktop\新建文件夹\GL\dl.xlsx"
    xg(datapath, dlpath)